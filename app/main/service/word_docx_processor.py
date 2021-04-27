import re
from pathlib import Path
from app.main.constant import SundayWorship
from datetime import date
from docx import Document
from app.main.log import logger


class WordParser:
    def __init__(self, input_file) -> None:
        file_path = (Path.cwd() / "downloads/" / input_file)
        if Path(file_path).is_file():
            pass
        else:
            print(f"file: {input_file} not found, use template file instead.")
            logger.info(f"file: {input_file} not found, use template file instead.")
            input_file = "2021夏季季表_測試衝突.docx"
            # input_file = "2020春季季表.docx"  # HINT for Debug

        # HINT link: https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205
        word = Document(Path.cwd() / "downloads/" / input_file)
        data = self.parse_docx(word)

        # self.tai_raw = data[2:14]
        # self.chi_raw = data[22:34]

        # HINT 這邊不能寫死，要判斷該行是否為輪值資料，以免遇到當月有三週或五週，就會抓錯資料
        self.tai_raw = []
        self.chi_raw = []
        index_l,  self.tai_raw = self.fetch_raw_duties(data, self.tai_raw)
        index = index_l[-1]
        index_l,  self.chi_raw = self.fetch_raw_duties(data[index:], self.chi_raw)

        self.chi_duties = dict()
        self.tai_duties = dict()
        self.sum_names = set()
        self.tai_index = {k: self.regex(v) for k, v in enumerate(data[1])}  # 原始值含有空白
        self.chi_index = {k: self.regex(v) for k, v in enumerate(data[21])}  # 原始值含有空白
        self.chi_subject = SundayWorship.chinese_subject
        self.tai_subject = SundayWorship.taiwan_subject

    def fetch_raw_duties(self, data, container: list):
        conv_flag = False
        index_l = []
        for index, elem in enumerate(data):
            if SundayWorship.month_convert.get(elem[0][0]):
                index_l.append(index)
                container.append(elem)
                conv_flag = True
            else:
                if conv_flag:
                    index_l.append(index)
                    return index_l, container

    def regex(self, s):
        return re.sub(r"[ \n\t]", "", s)

    def parse_docx(self, word):
        data = []
        # keys = None
        # table = word.tables[0]
        for table in word.tables:
            for i, row in enumerate(table.rows[1:]):
                text = (cell.text for cell in row.cells)
                row_data = row_data = tuple(text)
                data.append(row_data)
        return data

    def parsing_chinese_duty(self):
        for row in self.chi_raw:  # chinese
            people_duties = dict()
            month = SundayWorship.month_convert.get(self.regex(row[0])[0])
            day = int(self.regex(row[1]))
            title_date = date(2021, month, day)
            for index, name in enumerate(row):
                if index in (2, 3, 4, 5, 6, 11):
                    name = self.regex(name)
                    if name is not None and (name != ''):  # 有None就不要收
                        self.sum_names.add(name)
                        people_duties.setdefault(name, []).append(self.chi_subject.get(self.chi_index.get(index)))
            self.chi_duties.update({title_date: people_duties})

    def parsing_taiwanese_duty(self):
        for row in self.tai_raw:
            people_duties = dict()
            month = SundayWorship.month_convert.get(self.regex(row[0])[0])
            day = int(self.regex(row[2]))
            title_date = date(2021, month, day)
            for index, name in enumerate(row):
                if index in (3, 4, 5, 7, 10):
                    name = self.regex(name)
                    if name is not None and (name != ''):  # 有None就不要收
                        self.sum_names.add(name)
                        people_duties.setdefault(name, []).append(self.tai_subject.get(self.tai_index.get(index)))
                elif index == 8:
                    money_getters = name.split(" ")
                    people_duties.setdefault(self.regex(money_getters[0]), []).append(
                        self.tai_subject.get(self.tai_index.get(index)))
                    people_duties.setdefault(self.regex(money_getters[1]), []).append(
                        self.tai_subject.get(self.tai_index.get(index)))
                    self.sum_names.add(self.regex(money_getters[0]))
                    self.sum_names.add(self.regex(money_getters[1]))
            self.tai_duties.update({title_date: people_duties})

    def parsing_church_schedule(self):
        self.parsing_chinese_duty()
        self.parsing_taiwanese_duty()

        chi_date = list(self.chi_duties.keys())
        tai_date = list(self.tai_duties.keys())
        sum_date = set(chi_date + tai_date)

        # HINT 將taiduties整併到chiduties，並return出去
        for each_date in sum_date:
            if self.chi_duties.get(each_date) is None:
                self.chi_duties.update({each_date: dict()})
            for each_name in self.sum_names:
                lis1 = self.chi_duties[each_date].get(each_name)
                lis2 = self.tai_duties[each_date].get(each_name)
                if lis1 is None and lis2 is not None:
                    self.chi_duties[each_date][each_name] = self.tai_duties[each_date][each_name]
                elif lis1 is None and lis2 is None:
                    continue
                elif lis1 is not None and lis2 is None:
                    continue
                else:
                    self.chi_duties[each_date][each_name] = lis1 + lis2
        return self.chi_duties


class PostProcess:
    def __init__(self) -> None:
        pass

    def check_duplicate(self, tasks: list):
        time_slots = []
        for task in tasks:
            time_slots += SundayWorship.subject_slot.get(task)
        seen = set()
        repeat = set()
        for x in time_slots:
            if x not in seen:
                seen.add(x)
            else:
                repeat.add(x)
        return repeat

    def _generate_message(self, msg_collection: dict) -> str:
        msg_content = ""
        if msg_collection:
            for name_with_date, v in msg_collection.items():
                multi_tasks = ""
                multi_hours = ""
                for tasks, time_slots in v.items():
                    multi_tasks += tasks
                    for slot in time_slots:
                        multi_hours += f"{SundayWorship.slot_time.get(slot)} "
                msg_content += (f"{name_with_date}的服事分配有潛在問題， {multi_tasks}的時間重疊了: {multi_hours}\n"
                                f"▪️▪️▪️▪️▪️▪️▪️▪️▪️▪️\n")
        else:
            msg_content = "檢查完畢，服事內容未發現潛在問題。"
        return msg_content

    def check_conflict(self, member_duties: dict):
        repeat_collection = dict()
        conflict_flag = 0
        repeat = None
        for title_date, value in member_duties.items():
            for name, tasks in value.items():
                if name in SundayWorship.contact_namelist:
                    pass
                else:
                    repeat = self.check_duplicate(tasks)
                if repeat:
                    task_marker = dict()
                    time_slots = sorted(list(repeat))
                    for elem in time_slots:  # HINT 從user特有重複的timeslot內，找含有那些tasks，iter次數不多，且只有repeat存在時才會執行
                        task_marker.update({elem: []})
                        for task in tasks:
                            if elem in SundayWorship.subject_slot.get(task):
                                task_marker[elem].append(task)
                    task_timeslot = dict()
                    for slot, tasks in task_marker.items():
                        task_ring = ""
                        for task in tasks:
                            task_ring += f"{task} "  # HINT 每個task後面加上空白，用於顯示到Line畫面時可明顯區隔
                        task_timeslot.setdefault(f"{title_date} {task_ring}", []).append(slot)
                    repeat_collection.update({name: task_timeslot})
                    conflict_flag = 1

        msg_content = self._generate_message(repeat_collection)
        return msg_content, conflict_flag
